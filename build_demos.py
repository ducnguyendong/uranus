import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

folder_20 = r"D:\Du an dich tan cuc dai toan\20"
demo_dir = r"C:\Users\Nguyendd\Desktop\demo"
md_path = os.path.join(folder_20, "dich_thuat.md")
img_path = os.path.join(folder_20, "page_20_board_1.png")

with open(md_path, 'r', encoding='utf-8') as f:
    content = f.read()

def create_demo_A():
    doc = docx.Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    for line in content.split('\n'):
        line = line.strip()
        if not line: continue
        if "[[HÌNH_1]]" in line:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(img_path, width=Inches(3.2))
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_cap.add_run("(Hình 11: Song Chốt khéo thắng Đơn Sĩ Tượng)")
            run.italic = True
            run.font.size = Pt(11)
        else:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.save(os.path.join(demo_dir, "Demo_A_Chuan_Chuyen_Gia.docx"))

def create_demo_B():
    doc = docx.Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    for line in content.split('\n'):
        line = line.strip()
        if not line: continue
        if "[[HÌNH_1]]" in line:
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_cap.add_run("HÌNH 11 - THẾ CÔNG CỦA SONG CHỐT")
            run.bold = True
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(img_path, width=Inches(3.8))
        else:
            p = doc.add_paragraph()
            run = p.add_run(line)
            if line[0].isdigit() or "①" in line: run.bold = True
    doc.save(os.path.join(demo_dir, "Demo_B_Tan_Thoi.docx"))

def create_demo_C():
    doc = docx.Document()
    style = doc.styles['Normal']
    style.font.name = 'Palatino Linotype'
    style.font.size = Pt(12.5)
    
    for line in content.split('\n'):
        line = line.strip()
        if not line: continue
        if "[[HÌNH_1]]" in line:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(img_path, width=Inches(3.0))
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_cap.add_run("【 Hình 11 】")
            run.bold = True
        else:
            p = doc.add_paragraph(f"    {line}") # Thụt lề thủ công
    doc.save(os.path.join(demo_dir, "Demo_C_Co_Pho.docx"))

if __name__ == "__main__":
    create_demo_A()
    create_demo_B()
    create_demo_C()
    print("Xong: 3 bản demo đã được tạo trên Desktop!")
