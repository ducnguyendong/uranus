
import os
import time
import fitz
import cv2
import numpy as np
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Cấu hình đường dẫn
BASE_DIR = r"D:\Du an dich tan cuc dai toan"
OUTPUT_DIR = os.path.join(BASE_DIR, "output")

def crop_board_v6_logic(pdf_path, output_img_path):
    """Sử dụng logic v6 đã thành công để cắt bàn cờ"""
    try:
        doc = fitz.open(pdf_path)
        page = doc[0]
        pix = page.get_pixmap(matrix=fitz.Matrix(3, 3))
        img_data = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.h, pix.w, pix.n)
        img = cv2.cvtColor(img_data, cv2.COLOR_RGB2BGR) if pix.n == 3 else cv2.cvtColor(img_data, cv2.COLOR_RGBA2BGR)
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        _, thresh = cv2.threshold(gray, 200, 255, cv2.THRESH_BINARY_INV)
        contours, _ = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        
        best_box = None
        max_area = 0
        for cnt in contours:
            x, y, w, h = cv2.boundingRect(cnt)
            area = w * h
            aspect_ratio = float(w)/h
            if area > 100000 and 0.7 < aspect_ratio < 1.3:
                if area > max_area:
                    max_area = area
                    best_box = (x, y, w, h)
        
        if best_box:
            x, y, w, h = best_box
            res = img[y:y+h, x:x+w]
            cv2.imwrite(output_img_path, res)
            return True
    except Exception as e:
        print(f"Lỗi cắt ảnh: {e}")
    return False

def create_docx_template(output_docx, image_path, content_data):
    """Tạo file Word theo layout v3 chuẩn"""
    doc = Document()
    # (Phần này sẽ được tích hợp nội dung dịch từ AI)
    # Tạm thời để placeholder để test quy trình chuyển đổi
    doc.add_heading(content_data['title'], 0)
    table = doc.add_table(rows=1, cols=2)
    left_cell = table.cell(0, 0)
    p_img = left_cell.paragraphs[0]
    p_img.add_run().add_picture(image_path, width=Inches(2.5))
    left_cell.add_paragraph(content_data['caption']).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    right_cell = table.cell(0, 1)
    right_cell.add_paragraph(content_data['moves'])
    doc.save(output_docx)

def process_page(page_num):
    pdf_path = os.path.join(BASE_DIR, f"{page_num}.pdf")
    page_output_dir = os.path.join(OUTPUT_DIR, str(page_num))
    if not os.path.exists(page_output_dir):
        os.makedirs(page_output_dir)
        
    img_path = os.path.join(page_output_dir, f"board_{page_num}.png")
    docx_path = os.path.join(page_output_dir, f"{page_num}.docx")
    
    print(f"--- Đang xử lý trang {page_num} ---")
    
    # 1. Cắt ảnh
    if crop_board_v6_logic(pdf_path, img_path):
        print(f"  [v] Đã cắt bàn cờ.")
    else:
        print(f"  [x] Không tìm thấy bàn cờ trang {page_num}.")
        
    # 2. Bước dịch thuật (Sẽ gọi qua OpenClaw hoặc script AI)
    # ... logic dịch ...
    
    print(f"  [v] Hoàn thành trang {page_num}. Nghỉ 3 giây...")
    time.sleep(3)

if __name__ == "__main__":
    # Test thử trang 14 đã có
    process_page(14)
