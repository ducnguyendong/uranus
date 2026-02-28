
import fitz
import cv2
import numpy as np
import os
import json
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- CONFIG ---
WORKSPACE_DIR = r"D:\Du an dich tan cuc dai toan"
TEMP_DIR = os.path.join(WORKSPACE_DIR, "temp")
OUTPUT_DIR = os.path.join(WORKSPACE_DIR, "output")

if not os.path.exists(TEMP_DIR): os.makedirs(TEMP_DIR)
if not os.path.exists(OUTPUT_DIR): os.makedirs(OUTPUT_DIR)

class UranusMaster:
    def __init__(self, page_num):
        self.page_num = page_num
        self.pdf_path = os.path.join(WORKSPACE_DIR, f"{page_num}.pdf")
        self.full_page_img = os.path.join(TEMP_DIR, f"page_{page_num}_full.png")
        self.board_images = []

    def prepare_assets(self):
        """Chụp ảnh toàn trang và cắt các bàn cờ"""
        doc = fitz.open(self.pdf_path)
        page = doc[0]
        # Render high-res
        pix = page.get_pixmap(matrix=fitz.Matrix(3, 3))
        img_data = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.h, pix.w, pix.n)
        img = cv2.cvtColor(img_data, cv2.COLOR_RGB2BGR) if pix.n == 3 else cv2.cvtColor(img_data, cv2.COLOR_RGBA2BGR)
        
        # Tạo thư mục riêng cho trang
        page_dir = os.path.join(WORKSPACE_DIR, str(self.page_num))
        if not os.path.exists(page_dir): os.makedirs(page_dir)
        
        # Lưu ảnh toàn trang cho AI soi
        self.full_page_img = os.path.join(page_dir, f"page_{self.page_num}_full.png")
        cv2.imwrite(self.full_page_img, img)
        
        # Tìm và cắt bàn cờ (Logic v6)
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        _, thresh = cv2.threshold(gray, 200, 255, cv2.THRESH_BINARY_INV)
        contours, _ = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        
        raw_boards = []
        for cnt in contours:
            x, y, w, h = cv2.boundingRect(cnt)
            if (w*h) > 80000 and 0.7 < float(w)/h < 1.3:
                raw_boards.append((x, y, w, h))
        
        raw_boards.sort(key=lambda b: b[1]) # Sắp xếp từ trên xuống
        
        for i, (x, y, w, h) in enumerate(raw_boards):
            board_path = os.path.join(page_dir, f"page_{self.page_num}_board_{i+1}.png")
            cv2.imwrite(board_path, img[y:y+h, x:x+w])
            self.board_images.append(board_path)
        
        return self.full_page_img, self.board_images

    def build_docx(self, content_json_path):
        """Xây dựng file Word từ dữ liệu AI cung cấp"""
        with open(content_json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            
        doc = Document()
        # Setup margins
        for section in doc.sections:
            section.top_margin = Inches(0.6); section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(1.0); section.right_margin = Inches(1.0)

        # Header
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(data.get('header', "Tượng Kỳ Tàn Cục Đại Toàn"))
        run.italic = True; run.font.size = Pt(10); run.font.name = 'Times New Roman'

        # Blocks
        for block in data['blocks']:
            if block['type'] == 'text':
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for part in block['content']:
                    run = p.add_run(part['text'])
                    run.bold = part.get('bold', False)
                    run.italic = part.get('italic', False)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(block.get('size', 12))
            
            elif block['type'] == 'chess_layout':
                table = doc.add_table(rows=1, cols=2)
                table.autofit = False
                table.columns[0].width = Inches(3.0)
                table.columns[1].width = Inches(3.7)
                
                # Left: Image
                idx = block['board_index'] - 1
                c_left = table.cell(0, 0)
                p_img = c_left.paragraphs[0]
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.add_run().add_picture(self.board_images[idx], width=Inches(2.5))
                
                p_cap = c_left.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p_cap.add_run(block['figure_label'])
                run.font.size = Pt(10); run.font.name = 'Times New Roman'
                
                # Right: Moves
                c_right = table.cell(0, 1)
                for move_set in block['moves']:
                    pm = c_right.add_paragraph()
                    run_m = pm.add_run(move_set[0])
                    run_m.bold = True; run_m.font.name = 'Times New Roman'; run_m.font.size = Pt(11)
                    if len(move_set) > 1 and move_set[1]:
                        run_n = pm.add_run(f"\n{move_set[1]}")
                        run_n.italic = True; run_n.font.name = 'Times New Roman'; run_n.font.size = Pt(10)

        save_path = os.path.join(WORKSPACE_DIR, str(self.page_num), f"Trang_{self.page_num}.docx")
        doc.save(save_path)
        return save_path

if __name__ == "__main__":
    import sys
    mode = sys.argv[1] # prep / build
    page = sys.argv[2]
    master = UranusMaster(page)
    if mode == "prep":
        full, boards = master.prepare_assets()
        print(f"PREP_DONE|{full}|{len(boards)}")
    elif mode == "build":
        page_dir = os.path.join(WORKSPACE_DIR, str(page))
        # Count boards if not provided
        if len(sys.argv) > 4:
            b_count = int(sys.argv[4])
        else:
            b_count = len([f for f in os.listdir(page_dir) if f.startswith(f"page_{page}_board_")])
        master.board_images = [os.path.join(page_dir, f"page_{page}_board_{i+1}.png") for i in range(b_count)]
        path = master.build_docx(sys.argv[3])
        print(f"BUILD_DONE|{path}")
