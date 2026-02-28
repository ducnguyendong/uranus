
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_styled_run(paragraph, text, bold=False, italic=False, font_size=12):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(font_size)
    return run

def build_docx_16(output_path, image_path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Header trang
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_styled_run(p, "Tượng Kỳ Tàn Cục Đại Toàn", italic=True, font_size=10)

    # --- VÍ DỤ 4 ---
    p = doc.add_paragraph()
    add_styled_run(p, "Ví dụ 4", bold=True, font_size=14)

    # Đoạn dẫn nhập
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_styled_run(p, "Như (Hình 4). Đây là một ví dụ khác về Đơn Chốt phối hợp với Tướng khéo thắng. Đỏ dùng Tướng chiếm lộ sườn, Chốt khống chế trung lộ.")

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(3.0)
    table.columns[1].width = Inches(3.7)
    
    # Trái: Bàn cờ 4
    left_cell = table.cell(0, 0)
    p_img = left_cell.paragraphs[0]
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(image_path, width=Inches(2.5))
    add_styled_run(left_cell.add_paragraph(), "(Hình 4)", font_size=10).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Phải: Nội dung các nước đi
    right_cell = table.cell(0, 1)
    moves = [
        ("1. Tướng 5 bình 4", "Tướng 5 tiến 1"),
        ("2. Chốt 4 tiến 1", "Tướng 5 bình 6"),
        ("3. Tướng 4 bình 5", "(Đỏ thắng)"),
        ("", "Ghi chú: Ở thế cờ này, việc Đỏ chiếm được lộ 4 là cực kỳ quan trọng để phong tỏa hướng di chuyển của Hắc Tướng.")
    ]
    
    for i, (m, n) in enumerate(moves):
        pm = right_cell.add_paragraph()
        if m: add_styled_run(pm, m, bold=True, font_size=11)
        if n: add_styled_run(pm, f"\n{n}" if m else n, italic=True if not m else False, font_size=10)

    # Phần văn bản cuối trang (nếu có)
    doc.add_paragraph()
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_styled_run(p_footer, "Tóm lại, Đơn Chốt tuy năng lực có hạn nhưng nếu biết phối hợp chặt chẽ với Tướng để chiếm các lộ trọng yếu, vẫn có thể tạo nên những chiến thắng khéo léo trước Tướng đơn độc.")

    doc.save(output_path)

# CHẠY CHO TRANG 16
image16 = r"D:\Du an dich tan cuc dai toan\output\16\board_16_1.png"
build_docx_16(r"D:\Du an dich tan cuc dai toan\output\16\16_full_content.docx", image16)
print("DONE_16")
