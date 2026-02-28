from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx():
    doc = Document()
    
    # Thiết lập font chữ chung
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # --- Ví dụ 5 ---
    p5 = doc.add_paragraph()
    run5 = p5.add_run('Ví dụ 5')
    run5.bold = True
    p5.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph('Như (Hình 21), bên Đỏ có Sĩ, Tốt Đen không thể khống chế được Tướng Đỏ, Song Binh có thể ung dung thắng Tốt Sĩ.')

    # Chèn Hình 21
    try:
        p_img21 = doc.add_paragraph()
        p_img21.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img21 = p_img21.add_run()
        run_img21.add_picture(r'D:\Du an dich tan cuc dai toan\26\hinh_21.png', width=Inches(3.0))
        
        p_cap21 = doc.add_paragraph()
        p_cap21.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cap21 = p_cap21.add_run('(Hình 21)')
        run_cap21.italic = True
    except:
        doc.add_paragraph('[Thiếu Hình 21]')

    doc.add_paragraph('1. Binh năm tiến một')
    doc.add_paragraph('Nếu đi Tướng 5 thoái 6, Binh sáu tiến một, Sĩ 5 bình 4, Tướng năm bình sáu, Tướng 4 bình 4 (sic - có lẽ là Tướng 4 bình 5), Tướng sáu tiến một, Tướng 4 tiến 1, Binh sáu thoái một, Tướng 4 bình 5, Binh sáu tiến một sát, Đỏ thắng.')
    
    doc.add_paragraph('2. Tướng năm bình bốn')
    doc.add_paragraph('3. Binh sáu bình bảy')
    doc.add_paragraph('4. ...')
    
    doc.add_paragraph('Nếu đổi đi Sĩ 6 tiến 1, Binh bảy tiến một, Tướng 5 thoái 4, Binh bảy bình sáu, Tướng 5 bình 6, Tướng năm bình bốn...')
    
    doc.add_paragraph('5. Binh bảy bình sáu ...')
    doc.add_paragraph('Đến đây, bên Đen có hai cách đối phó, lần lượt trình bày như sau:')
    
    doc.add_paragraph('(Một)', style='List Bullet')
    doc.add_paragraph('5. ... Tướng 4 bình 5')
    doc.add_paragraph('6. Binh sáu tiến một, Tướng 5 thoái 4')
    doc.add_paragraph('7. Tướng năm bình bốn, Tướng 4 tiến 5')
    doc.add_paragraph('8. ...')
    doc.add_paragraph('9. Tướng bốn bình năm, Tướng 6 bình 5')
    doc.add_paragraph('10. Tướng năm bình sáu, Tướng 5 bình 4')
    doc.add_paragraph('11. Tướng sáu tiến bốn, Tướng 4 tiến 1')
    doc.add_paragraph('12. Binh bốn thoái năm, Tướng 6 tiến 5')
    doc.add_paragraph('13. Binh sáu bình năm, Tướng 5 bình 4')
    doc.add_paragraph('14. Tướng bốn tiến một (Đỏ thắng)')

    doc.add_paragraph('(Hai)', style='List Bullet')
    doc.add_paragraph('5. ...')
    doc.add_paragraph('7. Tướng bốn tiến một, Tướng 6 tiến 1')
    doc.add_paragraph('8. Tướng bốn thoái một, Tướng 6 thoái 1')
    doc.add_paragraph('9. Tướng năm bình bốn, Tướng 6 bình 7')
    doc.add_paragraph('10. Tướng bốn tiến một, Tướng 6 bình 5')
    doc.add_paragraph('11. Tướng bốn tiến một (Đỏ thắng)')

    # --- Ví dụ 6 ---
    p6 = doc.add_paragraph()
    run6 = p6.add_run('Ví dụ 6')
    run6.bold = True
    p6.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph('Song Binh đối Tốt Tượng cũng có thể tạo ra đòn tấn công mạnh mẽ, như (Hình 22) là một ví dụ đơn giản.')

    # Chèn Hình 22
    try:
        p_img22 = doc.add_paragraph()
        p_img22.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img22 = p_img22.add_run()
        run_img22.add_picture(r'D:\Du an dich tan cuc dai toan\26\hinh_22.png', width=Inches(3.0))
        
        p_cap22 = doc.add_paragraph()
        p_cap22.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cap22 = p_cap22.add_run('(Hình 22)')
        run_cap22.italic = True
    except:
        doc.add_paragraph('[Thiếu Hình 22]')

    doc.save(r'D:\Du an dich tan cuc dai toan\26\Trang_26_Hoan_Chinh_v2.docx')

create_docx()
