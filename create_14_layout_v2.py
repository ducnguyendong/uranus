
import os
from docx import Document
from docx.shared import Inches, Pt, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import qn
from docx.oxml import OxmlElement

def create_chess_docx_v2(output_path, image_path):
    doc = Document()

    # Thiết lập lề trang
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # 1. Chương 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CHƯƠNG 1: TÀN CUỘC LOẠI BINH ĐƠN GIẢN")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'

    # 2. Đoạn văn dẫn nhập
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run("Trong các quân tấn công, ").font.name = 'Times New Roman'
    run = p.add_run("Binh")
    run.bold = True
    run.font.name = 'Times New Roman'
    p.add_run(" có uy lực nhỏ nhất, cách đi đơn giản, sau khi qua sông mới có thể đi ngang. Tàn cuộc loại ").font.name = 'Times New Roman'
    run = p.add_run("Binh")
    run.bold = True
    run.font.name = 'Times New Roman'
    p.add_run(" nhìn chung tương đối đơn giản, dễ nắm bắt. Nay bắt đầu bàn từ Đơn ").font.name = 'Times New Roman'
    run = p.add_run("Binh")
    run.bold = True
    run.font.name = 'Times New Roman'
    p.add_run(" và Song ").font.name = 'Times New Roman'
    run = p.add_run("Binh.")
    run.bold = True
    run.font.name = 'Times New Roman'

    # 3. Tiết 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TIẾT 1: ĐƠN BINH KHẾO THẮNG")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    # 4. Nội dung tiết 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run("Năng lực của ").font.name = 'Times New Roman'
    run = p.add_run("Binh")
    run.bold = True
    run.font.name = 'Times New Roman'
    p.add_run(" có hạn, Đơn ").font.name = 'Times New Roman'
    run = p.add_run("Binh")
    run.bold = True
    run.font.name = 'Times New Roman'
    p.add_run(" chỉ có thể thắng ").font.name = 'Times New Roman'
    run = p.add_run("Tướng")
    run.bold = True
    run.font.name = 'Times New Roman'
    p.add_run(" đơn độc. Trong những tình huống đặc biệt, ").font.name = 'Times New Roman'
    run = p.add_run("Binh")
    run.bold = True
    run.font.name = 'Times New Roman'
    p.add_run(" phối hợp với ").font.name = 'Times New Roman'
    run = p.add_run("Tướng, Sĩ, Tượng").font.name = 'Times New Roman'
    run.bold = True
    p.add_run(" có thể tạo thành một số cục diện khéo thắng.").font.name = 'Times New Roman'

    # 5. Ví dụ 1
    p = doc.add_paragraph()
    run = p.add_run("Ví dụ 1:")
    run.bold = True
    run.font.name = 'Times New Roman'

    # BỐ CỤC PDF GỐC: Bàn cờ bên TRÁI, Nước đi bên PHẢI (như trong ảnh anh vừa gửi)
    # Em dùng Table 1 hàng, 2 cột để cố định vị trí
    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Inches(3.2)
    table.columns[1].width = Inches(3.5)
    
    # --- CỘT TRÁI: Bàn cờ và Chú thích ---
    left_cell = table.cell(0, 0)
    
    # Chèn ảnh bàn cờ vào cột trái
    p_img = left_cell.paragraphs[0]
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img = p_img.add_run()
    run_img.add_picture(image_path, width=Inches(2.8))
    
    # Chú thích hình (Hình 1) dưới ảnh
    p_cap = left_cell.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cap = p_cap.add_run("(Hình 1)")
    run_cap.font.size = Pt(10)
    run_cap.font.name = 'Times New Roman'

    # --- CỘT PHẢI: Nước đi và giải thích ---
    right_cell = table.cell(0, 1)
    
    # Đoạn text dẫn dụ bên phải ảnh
    p_intro = right_cell.paragraphs[0]
    p_intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_intro.add_run("Như (Hình 1), là ví dụ Đơn ").font.name = 'Times New Roman'
    run = p_intro.add_run("Binh")
    run.bold = True
    p_intro.add_run(" phối hợp với ").font.name = 'Times New Roman'
    run = p_intro.add_run("Tướng").font.name = 'Times New Roman'
    run.bold = True
    p_intro.add_run(" khéo thắng.").font.name = 'Times New Roman'

    # Danh sách các nước đi
    moves = [
        ("1. Tướng 5 bình 4!", "(Nước cờ kiềm chế rất hay)"),
        ("... Tướng 6 tiến 1", ""),
        ("2. Binh 5 bình 6 - Tướng 6 bình 5", ""),
        ("3. Tướng 4 tiến 1 - Tướng 5 thoái 1", ""),
        ("4. Binh 6 tiến 1 - Sĩ 6 thoái 5", ""),
        ("5. Tướng 4 bình 5 (Đỏ thắng)", "")
    ]
    
    for main_move, note in moves:
        pm = right_cell.add_paragraph()
        run_m = pm.add_run(main_move)
        run_m.bold = True
        run_m.font.name = 'Times New Roman'
        run_m.font.size = Pt(11)
        if note:
            run_n = pm.add_run(f"\n{note}")
            run_n.italic = True
            run_n.font.name = 'Times New Roman'
            run_n.font.size = Pt(10)

    doc.save(output_path)

output_docx = r"D:\Du an dich tan cuc dai toan\output\14\14_final_layout.docx"
image_board = r"D:\Du an dich tan cuc dai toan\output\14\board_v6_final.png"

create_chess_docx_v2(output_docx, image_board)
print("DOCX_LAYOUT_DONE")
