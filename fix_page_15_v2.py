
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_styled_run(paragraph, text, bold=False, italic=False, font_size=12):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(font_size)
    return run

def build_docx_15(output_path, boards):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Header trang
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_styled_run(p, "2       Tượng Kỳ Tàn Cục Đại Toàn", italic=True, font_size=10)

    # --- VÍ DỤ 2 ---
    p = doc.add_paragraph()
    add_styled_run(p, "Ví dụ 2", bold=True, font_size=14)

    table1 = doc.add_table(rows=1, cols=2)
    table1.columns[0].width = Inches(3.0)
    table1.columns[1].width = Inches(3.7)
    
    # Trái: Bàn cờ 2
    c1_left = table1.cell(0, 0)
    p_img1 = c1_left.paragraphs[0]
    p_img1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img1.add_run().add_picture(boards[0], width=Inches(2.5))
    add_styled_run(c1_left.add_paragraph(), "(Hình 2)", font_size=10).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Phải: Nội dung ví dụ 2
    c1_right = table1.cell(0, 1)
    p_txt1 = c1_right.paragraphs[0]
    add_styled_run(p_txt1, "Như (Hình 2). Đỏ dùng Tướng khống chế Tướng, dùng Sĩ khéo thắng.")
    moves2 = [
        ("1. Chốt 4 bình 5", "Tướng 5 thoái 1"),
        ("2. Sĩ 4 tiến 5", "Tướng 5 thoái 1"),
        ("3. Chốt 5 tiến 1", "Tướng 5 bình 4"),
        ("4. Chốt 5 tiến 1", "(Đỏ thắng)")
    ]
    for m, n in moves2:
        pm = c1_right.add_paragraph()
        add_styled_run(pm, m, bold=True, font_size=11)
        if n: add_styled_run(pm, f" {n}", italic=True, font_size=10)

    # --- VÍ DỤ 3 ---
    doc.add_paragraph() # Khoảng cách
    p = doc.add_paragraph()
    add_styled_run(p, "Ví dụ 3", bold=True, font_size=14)

    # Đoạn dẫn nhập trước hình 3
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_styled_run(p, "Trong tình huống bình thường, chỉ Sĩ có thể thủ hòa Đơn Chốt. Như (Hình 3) là trường hợp đặc biệt, Đỏ dùng Chốt và Tướng khéo léo khống chế Hắc Tướng, tạo thành thế công kẹp từ hai hướng, tức \"tả hữu giáp công\", từ đó giành thắng lợi.")

    table2 = doc.add_table(rows=1, cols=2)
    # Trái: Bàn cờ 3
    c2_left = table2.cell(0, 0)
    p_img2 = c2_left.paragraphs[0]
    p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img2.add_run().add_picture(boards[1], width=Inches(2.5))
    add_styled_run(c2_left.add_paragraph(), "(Hình 3)", font_size=10).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Phải: Nội dung ví dụ 3
    c2_right = table2.cell(0, 1)
    moves3 = [
        ("1. Chốt 6 tiến 1!", "............"),
        ("", "Nếu đổi đi Chốt 6 bình 5? Sĩ 4 tiến 5, Chốt 5 tiến 1 - Sĩ 5 thoái 6! Chốt 5 tiến 1, Tướng 5 bình 4, Sĩ 4 bình 5 - Chốt 6 tiến 5, Chốt 6 bình 5 - Sĩ 5 thoái 6, Hòa cờ."),
        ("1. ............ ", "Sĩ 4 tiến 5"),
        ("", "Nếu đổi đi Tướng 5 tiến 1, Sĩ 4 tiến 5, Tướng 5 thoái 1, Chốt 6 tiến 1 - Sĩ 4 tiến 5, Sĩ 4 bình 5 Đỏ thắng."),
        ("2. Chốt 6 tiến 1", "Sĩ 5 tiến 6"),
        ("3. Sĩ 4 tiến 1!", "Sĩ 6 thoái 5"),
        ("4. Chốt 6 bình 5", "Tướng 5 bình 6"),
        ("5. Sĩ 4 bình 5", "(Đỏ thắng)")
    ]
    for m, n in moves3:
        pm = c2_right.add_paragraph()
        if m: add_styled_run(pm, m, bold=True, font_size=11)
        if n: add_styled_run(pm, f" {n}", italic=True, font_size=10)

    doc.save(output_path)

# CHẠY CHO TRANG 15
boards15 = [r"D:\Du an dich tan cuc dai toan\output\15\board_15_1.png", r"D:\Du an dich tan cuc dai toan\output\15\board_15_2.png"]
build_docx_15(r"D:\Du an dich tan cuc dai toan\output\15\15_full_content_v2.docx", boards15)
print("FIX_15_DONE")
