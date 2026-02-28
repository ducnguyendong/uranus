
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_styled_run(paragraph, text, bold=False, italic=False, font_size=12):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(font_size)
    return run

def build_docx_16_perfect(output_path, image_path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_styled_run(p, "Tượng Kỳ Tàn Cục Đại Toàn", italic=True, font_size=10)

    # --- VÍ DỤ 4 ---
    p = doc.add_paragraph()
    add_styled_run(p, "Ví dụ 4", bold=True, font_size=14)

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(3.0)
    table.columns[1].width = Inches(3.7)
    
    # Cột trái: Hình 4
    c_left = table.cell(0, 0)
    p_img = c_left.paragraphs[0]
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(image_path, width=Inches(2.5))
    add_styled_run(c_left.add_paragraph(), "(Hình 4)", font_size=10).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Cột phải: Nội dung ví dụ 4
    c_right = table.cell(0, 1)
    p_desc = c_right.paragraphs[0]
    add_styled_run(p_desc, "Như (Hình 4) là cuộc cờ dùng Sĩ trợ công giành thắng lợi, quân Sĩ này ngoài tác dụng che chắn bảo vệ Tướng, còn có thể hỗ trợ Hồng Chốt khống chế Hắc Tướng.")
    
    moves = [
        ("① Tướng 5 bình 4!", "............"),
        ("", "Đây là một nước chờ cần thiết, khiến Hắc Tướng hết cách đối phó. Ghi nhớ, nếu đi sai thành Chốt 5 tiến 1? Tướng 5 thoái 1, Tướng 5 bình 4, Tướng 4 tiến 5, Tướng 4 bình 5, Tướng 5 thoái 1, Tướng 5 tiến 1, Tướng 5 bình 6, Chốt 5 bình 4, Tướng 4 tiến 5, Chốt 5 tiến 1 - Tướng 5 thoái 4, Hòa cờ."),
        ("① ............ ", "Tướng 5 thoái 6"),
        ("② Chốt 5 bình 6", "Tướng 5 tiến 1"),
        ("", "Như đổi đi Tướng 5 bình 4, Tướng 5 bình 6, Tướng 6 tiến 5, Chốt 6 tiến 1 - Tướng 4 bình 5, Soái 6 tiến 1 - Tướng 5 tiến 1, Soái 4 bình 5 Đỏ thắng."),
        ("③ Chốt 5 bình 4", "Tướng 5 bình 1"),
        ("④ Chốt 6 tiến 1", "Tướng 6 thoái 5"),
        ("⑤ Chốt 5 bình 3", "Tướng 5 thoái 6"),
        ("⑥ Chốt 4 tiến 1", "Tướng 6 tiến 5"),
        ("⑦ Chốt 4 bình 5", "Tướng 5 bình 6"),
        ("⑧ Chốt 6 bình 5", "(Đỏ thắng)")
    ]
    
    for m, n in moves:
        pm = c_right.add_paragraph()
        if m: add_styled_run(pm, m, bold=True, font_size=11)
        if n: add_styled_run(pm, f"\n{n}" if m else n, italic=not m, font_size=10)

    # --- TIẾT 2 ---
    doc.add_paragraph()
    p_t2 = doc.add_paragraph()
    p_t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_styled_run(p_t2, "TIẾT 2: ĐƠN CHỐT LỆ HÒA", bold=True, font_size=14)

    p_intro2 = doc.add_paragraph()
    p_intro2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_styled_run(p_intro2, "Một số cục thế tàn cuộc tất yếu dẫn đến hòa cờ, mang ý nghĩa điển hình, được gọi là tàn cuộc \"Lệ Hòa\". Chúng ta nên học tập và nắm vững các loại hình lệ hòa khác nhau, đây là những kiến thức thường thức cần thiết cho người chơi cờ. Dưới đây xin trình bày chi tiết về tàn cuộc Đơn Chốt Lệ Hòa.")

    doc.save(output_path)

# THI CÔNG
image16 = r"D:\Du an dich tan cuc dai toan\output\16\board_16_1.png"
build_docx_16_perfect(r"D:\Du an dich tan cuc dai toan\output\16\16_full_content_v2.docx", image16)
print("PERFECT_16_DONE")
