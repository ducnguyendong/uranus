
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_chess_docx(output_path, image_path):
    doc = Document()

    # Thiết lập lề trang (giả định A4)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # 1. Chương 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CHƯƠNG 1: TÀN CUỘC LOẠI BINH ĐƠN GIẢN")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'

    # 2. Đoạn văn dẫn nhập
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run("Trong các quân tấn công, ")
    run = p.add_run("Binh")
    run.bold = True
    run = p.add_run(" có uy lực nhỏ nhất, cách đi đơn giản, sau khi qua sông mới có thể đi ngang. Tàn cuộc loại ")
    run = p.add_run("Binh")
    run.bold = True
    run = p.add_run(" nhìn chung tương đối đơn giản, dễ nắm bắt. Nay bắt đầu bàn từ Đơn ")
    run = p.add_run("Binh")
    run.bold = True
    run = p.add_run(" và Song ")
    run = p.add_run("Binh.")
    run.bold = True
    p.style.font.name = 'Times New Roman'
    p.style.font.size = Pt(12)

    # 3. Tiết 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TIẾT 1: ĐƠN BINH KHẾO THẮNG")
    run.bold = True
    run.font.size = Pt(14)

    # 4. Nội dung tiết 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run("Năng lực của ")
    run = p.add_run("Binh")
    run.bold = True
    run = p.add_run(" có hạn, Đơn ")
    run = p.add_run("Binh")
    run.bold = True
    run = p.add_run(" chỉ có thể thắng ")
    run = p.add_run("Tướng")
    run.bold = True
    run = p.add_run(" đơn độc. Trong những tình huống đặc biệt, ")
    run = p.add_run("Binh")
    run.bold = True
    run = p.add_run(" phối hợp với ")
    run = p.add_run("Tướng, Sĩ, Tượng")
    run.bold = True
    run = p.add_run(" có thể tạo thành một số cục diện khéo thắng.")

    # 5. Ví dụ 1
    p = doc.add_paragraph()
    run = p.add_run("Ví dụ 1:")
    run.bold = True
    p = doc.add_paragraph("Như (Hình 1), là ví dụ Đơn ")
    run = p.add_run("Binh")
    run.bold = True
    run = p.add_run(" phối hợp với ")
    run = p.add_run("Tướng")
    run.bold = True
    run = p.add_run(" khéo thắng.")

    # 6. Chèn hình bàn cờ (Board v6 Final)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=Inches(3.5))

    # 7. Chú thích hình (CỠ CHỮ NHỎ NHƯ ANH DẶN)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("(Hình 1)")
    run.font.size = Pt(10) # Cỡ chữ 10 cho nhỏ gọn
    run.font.name = 'Times New Roman'

    # 8. Các nước đi
    moves = [
        ("1. Tướng 5 bình 4!", "(Nước cờ kiềm chế rất hay)"),
        ("... Tướng 6 tiến 1", ""),
        ("2. Binh 5 bình 6 - Tướng 6 bình 5", ""),
        ("3. Tướng 4 tiến 1 - Tướng 5 thoái 1", ""),
        ("4. Binh 6 tiến 1 - Sĩ 6 thoái 5", ""),
        ("5. Tướng 4 bình 5 (Đỏ thắng)", "")
    ]
    
    for main_move, note in moves:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(1.5) # Thụt lề cho đẹp
        run = p.add_run(main_move)
        run.bold = True
        if note:
            p.add_run(f"\n{note}").italic = True

    doc.save(output_path)

output_docx = r"D:\Du an dich tan cuc dai toan\output\14\14.docx"
image_board = r"D:\Du an dich tan cuc dai toan\output\14\board_v6_final.png"

create_chess_docx(output_docx, image_board)
print("DOCX_DONE")
