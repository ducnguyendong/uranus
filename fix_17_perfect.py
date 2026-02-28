
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_styled_run(paragraph, text, bold=False, italic=False, font_size=12):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(font_size)
    return run

def build_docx_17_perfect(output_path, boards):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5); section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(1); section.right_margin = Inches(1)

    # Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_styled_run(p, "4       Tượng Kỳ Tàn Cục Đại Toàn", italic=True, font_size=10)

    # --- VÍ DỤ 1 (HÌNH 5) ---
    p = doc.add_paragraph()
    add_styled_run(p, "Ví dụ 1", bold=True, font_size=14)

    table1 = doc.add_table(rows=1, cols=2)
    table1.columns[0].width = Inches(3.0); table1.columns[1].width = Inches(3.7)
    
    # Left: Hình 5
    c1_left = table1.cell(0, 0)
    p_img1 = c1_left.paragraphs[0]; p_img1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img1.add_run().add_picture(boards[0], width=Inches(2.5))
    add_styled_run(c1_left.add_paragraph(), "(Hình 5)", font_size=10).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Right: Nội dung ví dụ 1
    c1_right = table1.cell(0, 1)
    add_styled_run(c1_right.paragraphs[0], "Như (Hình 5), là thế cờ Chỉ một Sĩ giữ hòa Đơn Chốt. Chỉ cần chú ý không để Tướng Đỏ khống chế Tướng Đen, một quân Sĩ Đen là đủ để chống lại sự tấn công của Đơn Chốt Đỏ.")
    
    moves1 = [
        ("① Chốt 5 tiến 1", "Tướng 5 thoái 6"),
        ("", "Nếu đổi đi Tướng 4 bình 5? Soái 4 bình 5, Tướng 5 thoái 6, Soái 5 bình 6, Tướng 6 bình 5, Soái 5 bình 6, Tướng 5 bình 4, Soái 6 bình 5, Đỏ thắng."),
        ("② Soái 4 bình 5", "Tướng 4 tiến 1"),
        ("③ Sĩ 5 thoái 4", "Tướng 4 tiến 1"),
        ("④ Soái 5 bình 6", "Sĩ 6 tiến 5"),
        ("⑤ Soái 6 bình 5", "Sĩ 5 thoái 6"),
        ("⑥ Soái 5 tiến 1", "Tướng 4 tiến 1"),
        ("", "(Hòa cờ)")
    ]
    for m, n in moves1:
        pm = c1_right.add_paragraph()
        if m: add_styled_run(pm, m, bold=True, font_size=11)
        if n: add_styled_run(pm, f"\n{n}" if m else n, italic=not m, font_size=10)

    # --- VÍ DỤ 2 (HÌNH 6) ---
    doc.add_paragraph()
    p = doc.add_paragraph()
    add_styled_run(p, "Ví dụ 2", bold=True, font_size=14)

    table2 = doc.add_table(rows=1, cols=2)
    
    # Left: Hình 6
    c2_left = table2.cell(0, 0)
    p_img2 = c2_left.paragraphs[0]; p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img2.add_run().add_picture(boards[1], width=Inches(2.5))
    add_styled_run(c2_left.add_paragraph(), "(Hình 6)", font_size=10).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Right: Nội dung ví dụ 2
    c2_right = table2.cell(0, 1)
    add_styled_run(c2_right.paragraphs[0], "Như (Hình 6), Đỏ có thể đi nước chờ, chỉ cần chú ý sự kiềm chế của Tướng Đỏ đối với Tướng và Chốt của Đen, khiến Tướng Đen không thể phát huy uy lực, là có thể giữ hòa.")
    
    moves2 = [
        ("① Soái 8 bình 7", "Tướng 5 bình 4"),
        ("② Soái 5 bình 4", "Chốt 5 bình 6"),
        ("③ Soái 4 bình 5!", "............"),
        ("", "Tướng Đỏ nhất định phải kịp thời chiếm giữ trung lộ, mới có thể giữ hòa."),
        ("③ ............ ", "Chốt 6 bình 5"),
        ("④ Soái 5 bình 4", "Tướng 4 bình 5"),
        ("⑤ Soái 4 bình 5", "(Hòa cờ)")
    ]
    for m, n in moves2:
        pm = c2_right.add_paragraph()
        if m: add_styled_run(pm, m, bold=True, font_size=11)
        if n: add_styled_run(pm, f"\n{n}" if m else n, italic=not m, font_size=10)

    doc.save(output_path)

# THI CÔNG
b17 = [r"D:\Du an dich tan cuc dai toan\temp\page_17_board_1.png", r"D:\Du an dich tan cuc dai toan\temp\page_17_board_2.png"]
build_docx_17_perfect(r"D:\Du an dich tan cuc dai toan\output\17.docx", b17)
print("FIX_17_DONE")
