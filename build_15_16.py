
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_styled_run(paragraph, text, bold=False, italic=False, font_size=12):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(font_size)
    return run

def create_15_16(page_num, image_path, output_path, intro, moves_list):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(3.0)
    table.columns[1].width = Inches(3.7)
    
    # Left
    left_cell = table.cell(0, 0)
    p_img = left_cell.paragraphs[0]
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(image_path, width=Inches(2.5))
    p_cap = left_cell.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_styled_run(p_cap, f"(Hình {page_num - 13})", font_size=10)

    # Right
    right_cell = table.cell(0, 1)
    p_intro = right_cell.paragraphs[0]
    p_intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_styled_run(p_intro, intro)
    
    for main_move, note in moves_list:
        pm = right_cell.add_paragraph()
        add_styled_run(pm, main_move, bold=True, font_size=11)
        if note:
            add_styled_run(pm, f"\n{note}", italic=True, font_size=10)

    doc.save(output_path)

# DATA TRANG 15
create_15_16(15, 
    r"D:\Du an dich tan cuc dai toan\output\15\board_15.png",
    r"D:\Du an dich tan cuc dai toan\output\15\15.docx",
    "Như (Hình 2), là ví dụ Đơn Chốt phối hợp với Tướng khéo thắng.",
    [
        ("1. Tướng 4 tiến 1", ""),
        ("... Tướng 6 tiến 1", ""),
        ("2. Chốt 6 tiến 1 - Tướng 6 bình 5", ""),
        ("3. Tướng 4 bình 5", "(Đỏ thắng)")
    ]
)

# DATA TRANG 16
create_15_16(16, 
    r"D:\Du an dich tan cuc dai toan\output\16\board_16.png",
    r"D:\Du an dich tan cuc dai toan\output\16\16.docx",
    "Như (Hình 3), là ví dụ Đơn Chốt phối hợp với Tướng khéo thắng (biến hóa khác).",
    [
        ("1. Tướng 5 bình 4", ""),
        ("... Tướng 5 tiến 1", ""),
        ("2. Chốt 4 tiến 1 - Tướng 5 bình 6", ""),
        ("3. Tướng 4 bình 5", "(Đỏ thắng)")
    ]
)
print("DONE_15_16")
