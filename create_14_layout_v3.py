
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_chess_docx_v3(output_path, image_path):
    doc = Document()

    # Thiết lập lề trang
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Helper function to add styled text
    def add_styled_run(paragraph, text, bold=False, italic=False, font_size=12):
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = 'Times New Roman'
        run.font.size = Pt(font_size)
        return run

    # 1. Chương 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_styled_run(p, "CHƯƠNG 1: TÀN CUỘC LOẠI BINH ĐƠN GIẢN", bold=True, font_size=16)

    # 2. Đoạn văn dẫn nhập
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_styled_run(p, "Trong các quân tấn công, ")
    add_styled_run(p, "Chốt", bold=True)
    add_styled_run(p, " có uy lực nhỏ nhất, cách đi đơn giản, sau khi qua sông mới có thể đi ngang. Tàn cuộc loại ")
    add_styled_run(p, "Chốt", bold=True)
    add_styled_run(p, " nhìn chung tương đối đơn giản, dễ nắm bắt. Nay bắt đầu bàn từ Đơn ")
    add_styled_run(p, "Chốt", bold=True)
    add_styled_run(p, " và Song ")
    add_styled_run(p, "Chốt.", bold=True)

    # 3. Tiết 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_styled_run(p, "TIẾT 1: ĐƠN CHỐT KHẾO THẮNG", bold=True, font_size=14)

    # 4. Nội dung tiết 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_styled_run(p, "Năng lực của ")
    add_styled_run(p, "Chốt", bold=True)
    add_styled_run(p, " có hạn, Đơn ")
    add_styled_run(p, "Chốt", bold=True)
    add_styled_run(p, " chỉ có thể thắng ")
    add_styled_run(p, "Tướng", bold=True)
    add_styled_run(p, " đơn độc. Trong những tình huống đặc biệt, ")
    add_styled_run(p, "Chốt", bold=True)
    add_styled_run(p, " phối hợp với ")
    add_styled_run(p, "Tướng, Sĩ, Tượng", bold=True)
    add_styled_run(p, " có thể tạo thành một số cục diện khéo thắng.")

    # 5. Ví dụ 1
    p = doc.add_paragraph()
    add_styled_run(p, "Ví dụ 1:", bold=True)

    # BỐ CỤC PDF GỐC: Bàn cờ bên TRÁI, Nước đi bên PHẢI
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(3.0)
    table.columns[1].width = Inches(3.7)
    
    # --- CỘT TRÁI: Bàn cờ và Chú thích ---
    left_cell = table.cell(0, 0)
    p_img = left_cell.paragraphs[0]
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img = p_img.add_run()
    run_img.add_picture(image_path, width=Inches(2.5))
    
    p_cap = left_cell.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_styled_run(p_cap, "(Hình 1)", font_size=10)

    # --- CỘT PHẢI: Nước đi và giải thích ---
    right_cell = table.cell(0, 1)
    p_intro = right_cell.paragraphs[0]
    p_intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_styled_run(p_intro, "Như (Hình 1), là ví dụ Đơn ")
    add_styled_run(p_intro, "Chốt", bold=True)
    add_styled_run(p_intro, " phối hợp với ")
    add_styled_run(p_intro, "Tướng", bold=True)
    add_styled_run(p_intro, " khéo thắng.")

    moves = [
        ("1. Tướng 5 bình 4!", "(Nước cờ kiềm chế rất hay)"),
        ("... Tướng 6 tiến 1", ""),
        ("2. Chốt 5 bình 6 - Tướng 6 bình 5", ""),
        ("3. Tướng 4 tiến 1 - Tướng 5 thoái 1", ""),
        ("4. Chốt 6 tiến 1 - Sĩ 6 thoái 5", ""),
        ("5. Tướng 4 bình 5 (Đỏ thắng)", "")
    ]
    
    for main_move, note in moves:
        pm = right_cell.add_paragraph()
        add_styled_run(pm, main_move, bold=True, font_size=11)
        if note:
            add_styled_run(pm, f"\n{note}", italic=True, font_size=10)

    doc.save(output_path)

output_docx = r"D:\Du an dich tan cuc dai toan\output\14\14_final_layout_v3.docx"
image_board = r"D:\Du an dich tan cuc dai toan\output\14\board_v6_final.png"

create_chess_docx_v3(output_docx, image_board)
print("DOCX_LAYOUT_V3_DONE")
